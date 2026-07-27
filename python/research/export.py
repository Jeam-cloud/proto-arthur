"""Render a finished paper to .docx or .pdf.

Both renderers walk the SAME paper dict and use the SAME citation formatter,
so a Word file and a PDF of one paper are not two independent
interpretations that can drift -- they are two renderings of one document.

WHY the `[n]` markers are replaced here and not on screen: in the app a
citation is a control (hover to highlight the source, click to jump to it),
so it stays a numbered pill. In an exported file nobody can click anything,
so it becomes the actual citation the chosen style prescribes. Same data,
different medium, different correct answer.

WHY reportlab rather than weasyprint for PDF: weasyprint needs GTK/Pango
system libraries. On a Windows-first desktop app that turns "export a PDF"
into "install a C toolchain", which is not a trade worth making for a
document this simple. reportlab is pure Python and ships as a wheel.
"""

from __future__ import annotations

import io
import logging

from research import citations

log = logging.getLogger(__name__)


def _prepared(paper: dict, sources: list[dict], style: str) -> tuple[list[dict], dict[int, dict]]:
    """Sections with citations rendered into the prose, plus the lookup used
    to build the reference list. Shared by both renderers so they cannot
    disagree about what the text says.

    Blocks keep their kind: prose becomes a string, a table stays structured
    so each renderer can lay it out natively (a real Word table, a real PDF
    table) rather than flattening it to text.
    """
    by_n = {int(s["n"]): s for s in sources if s.get("n") is not None}
    out = []
    for sec in paper.get("sections", []):
        blocks = []
        for p in sec.get("paragraphs", []):
            if p.get("kind") == "table":
                blocks.append({"kind": "table", **p})
                continue
            text = citations.render_in_text(p.get("text", ""), by_n, style)
            blocks.append({"kind": "text", "text": citations.dedupe_adjacent(text)})
        out.append({"heading": sec.get("heading", ""), "blocks": blocks})
    return out, by_n


def _refs(paper: dict, sources: list[dict], style: str, prebuilt: list[dict] | None) -> list[str]:
    """Only sources the paper ACTUALLY cites go in the reference list.

    An evidence panel holds everything the investigation found; a reference
    list holds what the paper used. Padding the second with the first is a
    real form of academic dishonesty, so the filter is deliberate.
    """
    cited = {c for sec in paper.get("sections", []) for p in sec.get("paragraphs", []) for c in (p.get("citations") or [])}
    used = [s for s in sources if int(s.get("n") or 0) in cited]
    if prebuilt:
        keep = {e.get("n") for e in prebuilt if e.get("n") in cited}
        return [e["text"] for e in prebuilt if e.get("n") in keep]
    return [e["text"] for e in citations.reference_list(used, style)]


# ---------- Word ----------

def to_docx(paper: dict, sources: list[dict], style: str, prebuilt_refs: list[dict] | None = None) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    sections, _ = _prepared(paper, sources, style)
    doc = Document()

    # Double-spaced 12pt Times on 1in margins is the default every style guide
    # converges on, and it is what a marker expects to receive.
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(1)

    title = doc.add_paragraph(paper.get("title") or "Untitled")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True

    if paper.get("abstract"):
        h = doc.add_paragraph("Abstract")
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.runs[0].bold = True
        doc.add_paragraph(paper["abstract"])

    for sec in sections:
        h = doc.add_paragraph(sec["heading"])
        h.runs[0].bold = True
        for block in sec["blocks"]:
            if block["kind"] == "table":
                _docx_table(doc, block, style, Pt, Inches, WD_ALIGN_PARAGRAPH)
                continue
            p = doc.add_paragraph(block["text"])
            p.paragraph_format.first_line_indent = Inches(0.5)

    refs = _refs(paper, sources, style, prebuilt_refs)
    if refs:
        h = doc.add_paragraph(citations.HEADINGS.get(style, "References"))
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        h.runs[0].bold = True
        for entry in refs:
            p = doc.add_paragraph(entry)
            # Hanging indent: first line flush, continuations indented. Every
            # author-date style requires it and its absence is the single most
            # obvious tell of a bibliography that was not really formatted.
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_table(doc, block, style, Pt, Inches, WD_ALIGN_PARAGRAPH) -> None:
    """A real Word table, not an ASCII imitation -- the whole point of
    exporting to .docx is that the result is editable in Word."""
    cols = list(block.get("columns") or [])
    rows = list(block.get("rows") or [])
    srcs = list(block.get("row_sources") or [])
    if not cols or not rows:
        return

    table = doc.add_table(rows=1, cols=len(cols) + 1)
    table.style = "Table Grid"
    header = table.rows[0].cells
    for i, label in enumerate(cols):
        header[i].text = label
        for p in header[i].paragraphs:
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                run.bold = True
    header[len(cols)].text = "Src"
    for p in header[len(cols)].paragraphs:
        p.paragraph_format.line_spacing = 1.0
        for run in p.runs:
            run.bold = True

    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci, value in enumerate(row[:len(cols)]):
            cells[ci].text = str(value)
            for p in cells[ci].paragraphs:
                p.paragraph_format.line_spacing = 1.0
        n = srcs[ri] if ri < len(srcs) else ""
        # Src column stays a bare number in every style. In a table it is a
        # locator back to the reference list, not a running-text citation.
        cells[len(cols)].text = f"[{n}]" if n else ""
        for p in cells[len(cols)].paragraphs:
            p.paragraph_format.line_spacing = 1.0

    if block.get("caption"):
        cap = doc.add_paragraph(block["caption"])
        cap.paragraph_format.line_spacing = 1.0
        cap.paragraph_format.space_after = Pt(12)
        for run in cap.runs:
            run.italic = True


# ---------- PDF ----------

def to_pdf(paper: dict, sources: list[dict], style: str, prebuilt_refs: list[dict] | None = None) -> bytes:
    from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    sections, _ = _prepared(paper, sources, style)
    base = getSampleStyleSheet()

    body = ParagraphStyle(
        "Body", parent=base["Normal"], fontName="Times-Roman", fontSize=12,
        leading=24, firstLineIndent=0.5 * inch, alignment=TA_JUSTIFY, spaceAfter=0,
    )
    heading = ParagraphStyle(
        "Head", parent=base["Normal"], fontName="Times-Bold", fontSize=12,
        leading=24, spaceBefore=12, spaceAfter=0,
    )
    centered = ParagraphStyle(
        "Centered", parent=heading, alignment=TA_CENTER,
    )
    ref_style = ParagraphStyle(
        "Ref", parent=base["Normal"], fontName="Times-Roman", fontSize=12,
        leading=24, leftIndent=0.5 * inch, firstLineIndent=-0.5 * inch,
    )

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        topMargin=inch, bottomMargin=inch, leftMargin=inch, rightMargin=inch,
        title=paper.get("title") or "Paper",
    )

    flow = [Paragraph(_esc(paper.get("title") or "Untitled"), centered)]
    if paper.get("abstract"):
        flow += [Paragraph("Abstract", centered),
                 Paragraph(_esc(paper["abstract"]), ParagraphStyle("Abs", parent=body, firstLineIndent=0))]
    for sec in sections:
        flow.append(Paragraph(_esc(sec["heading"]), heading))
        for block in sec["blocks"]:
            if block["kind"] == "table":
                flow += _pdf_table(block, base, inch)
                continue
            flow.append(Paragraph(_esc(block["text"]), body))

    refs = _refs(paper, sources, style, prebuilt_refs)
    if refs:
        # References start on their own page: standard in APA and Chicago, and
        # harmless in the others.
        flow += [PageBreak(), Paragraph(citations.HEADINGS.get(style, "References"), centered), Spacer(1, 6)]
        flow += [Paragraph(_esc(r), ref_style) for r in refs]

    doc.build(flow, onFirstPage=_page_number, onLaterPages=_page_number)
    return buf.getvalue()


def _pdf_table(block: dict, base, inch) -> list:
    """reportlab flowables for one comparison table plus its caption."""
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle

    cols = list(block.get("columns") or [])
    rows = list(block.get("rows") or [])
    srcs = list(block.get("row_sources") or [])
    if not cols or not rows:
        return []

    data = [[*cols, "Src"]]
    for ri, row in enumerate(rows):
        n = srcs[ri] if ri < len(srcs) else ""
        data.append([*[str(c) for c in row[:len(cols)]], f"[{n}]" if n else ""])

    table = Table(data, hAlign="LEFT", repeatRows=1)
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "Times-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Times-Roman"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("LINEBELOW", (0, 0), (-1, 0), 0.75, colors.black),
        ("LINEABOVE", (0, 0), (-1, 0), 0.75, colors.black),
        ("LINEBELOW", (0, -1), (-1, -1), 0.75, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ALIGN", (-1, 0), (-1, -1), "RIGHT"),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))

    out = [Spacer(1, 10), table]
    if block.get("caption"):
        cap = ParagraphStyle("Cap", parent=base["Normal"], fontName="Times-Italic",
                             fontSize=10, leading=13, spaceBefore=6)
        out.append(Paragraph(_esc(block["caption"]), cap))
    out.append(Spacer(1, 12))
    return out


def _esc(text: str) -> str:
    """reportlab's Paragraph parses a mini-HTML dialect, so bare & < > in
    source titles would either vanish or raise. Escape before it gets there."""
    return (
        (text or "")
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _page_number(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFont("Times-Roman", 10)
    canvas.drawRightString(LETTER_WIDTH - 72, 40, str(doc.page))
    canvas.restoreState()


LETTER_WIDTH = 612  # points; reportlab LETTER is (612, 792)


def filename_for(paper: dict, ext: str) -> str:
    raw = (paper.get("title") or "investigation")[:60]
    safe = "".join(ch if ch.isalnum() or ch in " -_" else "" for ch in raw).strip()
    return (safe.replace(" ", "-") or "paper") + f".{ext}"
